		import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
import datetime

# --- List of ICAI Auditing Standards (extend as needed) ---
AUDITING_STANDARDS = {
    "SA 200": "Overall Objectives of the Independent Auditor and the Conduct of an Audit in Accordance with Standards on Auditing",
    "SA 210": "Agreeing the Terms of Audit Engagements",
    "SA 220": "Quality Control for an Audit of Financial Statements",
    "SA 230": "Audit Documentation",
    "SA 240": "The Auditor’s Responsibilities Relating to Fraud in an Audit of Financial Statements",
    "SA 250": "Consideration of Laws and Regulations in an Audit of Financial Statements",
    "SA 260": "Communication with Those Charged with Governance",
    "SA 265": "Communicating Deficiencies in Internal Control to Those Charged with Governance and Management",
    "SA 299": "Responsibility of Joint Auditors",
    "SA 300": "Planning an Audit of Financial Statements",
    "SA 315": "Identifying and Assessing the Risks of Material Misstatement",
    "SA 320": "Materiality in Planning and Performing an Audit",
    "SA 330": "Auditor’s Responses to Assessed Risks",
    "SA 402": "Audit Considerations Relating to an Entity Using a Service Organization",
    "SA 450": "Evaluation of Misstatements Identified during the Audit",
    "SA 500": "Audit Evidence",
    "SA 501": "Audit Evidence – Specific Considerations for Selected Items",
    "SA 505": "External Confirmations",
    "SA 510": "Initial Audit Engagements – Opening Balances",
    "SA 520": "Analytical Procedures",
    "SA 530": "Audit Sampling",
    "SA 540": "Auditing Accounting Estimates, Including Fair Value Accounting Estimates, and Related Disclosures",
    "SA 550": "Related Parties",
    "SA 560": "Subsequent Events",
    "SA 570": "Going Concern",
    "SA 580": "Written Representations",
    "SA 600": "Using the Work of Another Auditor",
    "SA 610": "Using the Work of Internal Auditors",
    "SA 620": "Using the Work of an Auditor’s Expert",
    "SA 700": "Forming an Opinion and Reporting on Financial Statements",
    "SA 701": "Communicating Key Audit Matters in the Independent Auditor’s Report",
    "SA 705": "Modifications to the Opinion in the Independent Auditor’s Report",
    "SA 706": "Emphasis of Matter Paragraphs and Other Matter Paragraphs in the Independent Auditor’s Report",
    "SA 710": "Comparative Information – Corresponding Figures and Comparative Financial Statements",
    "SA 720": "The Auditor’s Responsibilities Relating to Other Information",
    "SA 800": "Special Considerations—Audits of Financial Statements Prepared in Accordance with Special Purpose Frameworks",
    "SA 805": "Special Considerations—Audits of Single Financial Statements and Specific Elements, Accounts or Items of a Financial Statement",
    "SA 810": "Engagements to Report on Summary Financial Statements",
    # Extend with new/revised standards as needed.
}

st.set_page_config(page_title="ICAI Audit Compliance App", layout="wide")
st.title("ICAI Audit Compliance Web App")
st.write("Track compliance for all ICAI Auditing Standards (SAs), add notes, and export your audit trail to Excel or Word.")

client_name = st.text_input("Client Name", "XYZ Ltd")
audit_date = st.date_input("Audit Date", datetime.date.today())

# --- Initialize Compliance State ---
if "compliance" not in st.session_state:
    st.session_state.compliance = {sa: {"done": False, "notes": ""} for sa in AUDITING_STANDARDS}

def reset_compliance():
    st.session_state.compliance = {sa: {"done": False, "notes": ""} for sa in AUDITING_STANDARDS}

# --- Add New Standard ---
st.sidebar.header("Add New ICAI Standard")
with st.sidebar.form("add_sa_form"):
    new_sa_code = st.text_input("SA Code (e.g., SA 900)")
    new_sa_name = st.text_input("Standard Name/Description")
    submitted = st.form_submit_button("Add Standard")
    if submitted and new_sa_code and new_sa_name:
        if new_sa_code not in AUDITING_STANDARDS:
            AUDITING_STANDARDS[new_sa_code] = new_sa_name
            st.session_state.compliance[new_sa_code] = {"done": False, "notes": ""}
            st.success(f"Added {new_sa_code}: {new_sa_name}")
        else:
            st.warning(f"{new_sa_code} already exists.")

st.sidebar.button("Reset All Compliance Data", on_click=reset_compliance)

# --- Checklist UI ---
st.subheader(f"Audit Checklist for {client_name} ({audit_date})")

for sa, desc in AUDITING_STANDARDS.items():
    col1, col2 = st.columns([1, 5])
    with col1:
        st.session_state.compliance[sa]["done"] = st.checkbox(f"{sa}", value=st.session_state.compliance[sa]["done"])
    with col2:
        st.markdown(f"**{desc}**")
        st.session_state.compliance[sa]["notes"] = st.text_area(
            f"Notes for {sa}", value=st.session_state.compliance[sa]["notes"], key=sa
        )

# --- Summary ---
st.markdown("---")
st.subheader("Summary Report")
done_count = sum(1 for d in st.session_state.compliance.values() if d["done"])
total_count = len(st.session_state.compliance)
st.write(f"Completed: {done_count}/{total_count} standards")

if st.button("Show Pending Standards"):
    st.write("**Pending SAs:**")
    for sa, d in st.session_state.compliance.items():
        if not d["done"]:
            st.write(f"- {sa}: {AUDITING_STANDARDS[sa]}")

# --- Excel Export ---
def create_excel():
    data = []
    for sa, details in st.session_state.compliance.items():
        data.append({
            "SA Code": sa,
            "Standard Name": AUDITING_STANDARDS[sa],
            "Compliant": "Yes" if details["done"] else "No",
            "Notes": details["notes"],
        })
    df = pd.DataFrame(data)
    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Audit Compliance")
        pd.DataFrame([{"Client Name": client_name, "Audit Date": audit_date}]).to_excel(writer, sheet_name="Info", index=False)
    output.seek(0)
    return output

if st.button("Export Compliance Report (Excel)"):
    excel_file = create_excel()
    st.download_button(
        label="Download Excel Report",
        data=excel_file,
        file_name="audit_compliance_report.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# --- Word Export ---
def create_word():
    doc = Document()
    doc.add_heading(f'Audit Compliance Report for {client_name}', 0)
    doc.add_paragraph(f'Audit Date: {audit_date}')
    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'SA Code'
    hdr_cells[1].text = 'Standard Name'
    hdr_cells[2].text = 'Compliant'
    hdr_cells[3].text = 'Notes'
    for sa, details in st.session_state.compliance.items():
        row_cells = table.add_row().cells
        row_cells[0].text = sa
        row_cells[1].text = AUDITING_STANDARDS[sa]
        row_cells[2].text = "Yes" if details["done"] else "No"
        row_cells[3].text = details["notes"]

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

if st.button("Export Compliance Report (Word)"):
    word_file = create_word()
    st.download_button(
        label="Download Word Report",
        data=word_file,
        file_name="audit_compliance_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.info("This is a prototype. For multi-user support, database integration, or detailed checklists for each SA, further development is required.")

# Requirements:
# pip install streamlit pandas python-docx xlsxwriter